from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as exc:
    raise SystemExit(
        "Package python-docx belum terpasang. Jalankan: pip install python-docx"
    ) from exc


PROJECT_ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = PROJECT_ROOT / "screenshots"
MEDIA_DIR = PROJECT_ROOT / "docx_extracted" / "word" / "media"
OUTPUT_DOCX = PROJECT_ROOT / "Smart Pelayanan Masyarakat kelompok5.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    for name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)]:
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    return doc


def add_centered(doc, text, size=12, bold=False, spacing_after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    return paragraph


def add_justified(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.add_run(text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.add_run(item)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def find_image(*candidates):
    for candidate in candidates:
        path = PROJECT_ROOT / candidate
        if path.exists():
            return path
    return None


def add_image_or_placeholder(doc, image_path, caption, width_cm=15):
    if image_path and image_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        set_cell_text(cell, "Tempat screenshot/gambar: " + caption)
        set_cell_shading(cell, "F2F2F2")

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.bold = True
    caption_run.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        set_cell_shading(header_cells[idx], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    return table


def add_cover(doc):
    logo = find_image("docx_extracted/word/media/image1.png", "screenshots/image1.png")
    if logo:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(logo), width=Cm(3.2))

    add_centered(doc, "SISTEM PELAYANAN MASYARAKAT", 18, True, 4)
    add_centered(doc, "(SmartPelayanan)", 16, True, 18)
    add_centered(doc, "LAPORAN TUGAS BESAR", 14, True, 8)
    add_centered(doc, "Mata Kuliah: IFB-202 Pemrograman Berbasis Objek", 12, False, 4)
    add_centered(doc, "Kelompok: 5", 12, False, 4)
    add_centered(doc, "Dosen: Dr. Uung Ungkawa, S.T., M.T.", 12, False, 24)

    add_centered(doc, "Disusun oleh:", 12, True, 4)
    add_centered(doc, "Kelompok 5", 12, False, 20)
    add_centered(doc, "Program Studi Teknik Informatika", 12, False, 4)
    add_centered(doc, "2026", 12, False, 4)
    doc.add_page_break()


def add_preface(doc):
    add_heading(doc, "KATA PENGANTAR", 1)
    add_justified(
        doc,
        "Puji syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa karena laporan "
        "Tugas Besar Pemrograman Berbasis Objek ini dapat disusun dengan baik. "
        "Laporan ini menjelaskan perancangan dan implementasi Sistem Pelayanan "
        "Masyarakat berbasis web dengan nama SmartPelayanan."
    )
    add_justified(
        doc,
        "SmartPelayanan dibuat untuk membantu masyarakat mengirim pengaduan, "
        "mengajukan layanan administrasi, memantau status, serta membantu admin "
        "mengelola data secara lebih rapi dan transparan."
    )
    add_justified(
        doc,
        "Kami menyadari laporan ini masih dapat dikembangkan lebih lanjut. Oleh "
        "karena itu, kritik dan saran yang membangun sangat diharapkan untuk "
        "penyempurnaan sistem dan laporan ini."
    )
    doc.add_page_break()


def add_table_of_contents(doc):
    add_heading(doc, "DAFTAR ISI", 1)
    chapters = [
        "BAB I PENDAHULUAN",
        "BAB II LANDASAN TEORI",
        "BAB III PERANCANGAN SISTEM",
        "BAB IV IMPLEMENTASI SISTEM",
        "BAB V DEPLOYMENT DAN CLOUDFLARE TUNNEL",
        "BAB VI PENGUJIAN SISTEM",
        "BAB VII PENUTUP",
        "PENJELASAN KODE, ALUR, DAN PENERAPAN OOP",
        "DAFTAR GAMBAR",
    ]
    for chapter in chapters:
        doc.add_paragraph(chapter, style="List Number")
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "BAB I PENDAHULUAN", 1)
    add_heading(doc, "1.1 Latar Belakang", 2)
    add_justified(
        doc,
        "Perkembangan teknologi informasi dan komunikasi memberikan dampak besar "
        "dalam berbagai bidang, termasuk pelayanan masyarakat. Sistem manual yang "
        "mengandalkan pencatatan kertas, kedatangan langsung ke kantor, dan "
        "komunikasi tidak terpusat sering menyebabkan proses pelayanan menjadi "
        "lambat, sulit dipantau, dan kurang transparan."
    )
    add_justified(
        doc,
        "SmartPelayanan dirancang sebagai sistem pelayanan masyarakat berbasis web "
        "untuk memudahkan pengaduan, pengajuan layanan administrasi, pemantauan "
        "status, serta pengelolaan data oleh admin. Sistem ini memanfaatkan Java, "
        "Spring Boot, Spring Data JPA, REST API, dan PostgreSQL."
    )
    add_heading(doc, "1.2 Rumusan Masalah", 2)
    add_bullets(doc, [
        "Bagaimana merancang sistem pelayanan masyarakat berbasis web?",
        "Bagaimana mengelola data pengaduan dan layanan administrasi?",
        "Bagaimana menyediakan fitur pendaftaran pelayanan secara online?",
        "Bagaimana menyajikan informasi status layanan kepada masyarakat?",
    ])
    add_heading(doc, "1.3 Tujuan Penelitian", 2)
    add_bullets(doc, [
        "Merancang sistem pelayanan masyarakat berbasis web.",
        "Mengelola data pengaduan dan layanan administrasi secara terintegrasi.",
        "Mengembangkan fitur pendaftaran pelayanan online.",
        "Menyediakan informasi status layanan secara efektif.",
    ])
    add_heading(doc, "1.4 Manfaat Penelitian", 2)
    add_bullets(doc, [
        "Manfaat bagi instansi: mempercepat pengelolaan pengaduan dan layanan.",
        "Manfaat bagi masyarakat: memudahkan pengajuan dan pemantauan status.",
        "Manfaat bagi pengembang: menerapkan konsep OOP pada aplikasi nyata.",
    ])


def add_chapter_2(doc):
    add_heading(doc, "BAB II LANDASAN TEORI", 1)
    sections = [
        ("2.1 Object-Oriented Programming (OOP)", "OOP adalah paradigma pemrograman yang memodelkan sistem sebagai kumpulan objek. Objek memiliki atribut dan method, sehingga kode menjadi lebih mudah dipahami, dikembangkan, dan dirawat."),
        ("2.2 Encapsulation", "Encapsulation membungkus data dan perilaku dalam class. Pada project ini, field entity dibuat private dan diakses melalui getter/setter, misalnya pada User, Pengaduan, dan LayananAdministrasi."),
        ("2.3 Inheritance", "Inheritance memungkinkan class mewarisi atribut dan method dari class lain. Pada project ini, entity seperti User, Pengaduan, LayananAdministrasi, dan FormField mewarisi BaseEntity untuk id, createdAt, dan updatedAt."),
        ("2.4 Polymorphism", "Polymorphism membuat objek dapat diperlakukan melalui tipe yang lebih umum. Dalam Spring, banyak repository memiliki bentuk operasi yang sama melalui JpaRepository, tetapi bekerja pada entity berbeda."),
        ("2.4.1 Overloading", "Overloading adalah method dengan nama sama tetapi parameter berbeda. Konsep ini dapat diterapkan ketika service menyediakan variasi pencarian data, misalnya berdasarkan status, kategori, atau user."),
        ("2.4.2 Overriding", "Overriding adalah penulisan ulang method dari parent/interface. Repository Spring Data JPA menerapkan perilaku dari interface JpaRepository untuk entity masing-masing."),
        ("2.5 Abstraction", "Abstraction menyembunyikan detail rumit dan menampilkan operasi penting. Controller cukup memanggil service/repository tanpa mengetahui detail query database."),
        ("2.5.1 Interface", "Interface digunakan pada repository seperti UserRepository, PengaduanRepository, dan LayananAdministrasiRepository. Spring membuat implementasi otomatis saat aplikasi berjalan."),
        ("2.5.2 Abstract Class", "BaseEntity berfungsi seperti rancangan umum untuk entity lain. Semua entity turunan mendapat id dan waktu pencatatan yang konsisten."),
        ("2.5.3 Exception Handling", "Exception handling digunakan dengan try-catch di controller untuk mengembalikan response ketika terjadi error seperti data tidak ditemukan atau request tidak valid."),
        ("2.6 Bahasa Pemrograman Java", "Java dipakai karena mendukung OOP dengan kuat, stabil untuk aplikasi backend, dan memiliki ekosistem Spring Boot yang matang."),
        ("2.7 Spring Boot", "Spring Boot memudahkan pembuatan aplikasi web Java melalui konfigurasi otomatis, embedded server, dependency injection, dan integrasi database."),
        ("2.8 Spring Data JPA", "Spring Data JPA memudahkan akses database menggunakan repository. Developer cukup membuat interface, lalu Spring menangani operasi CRUD."),
        ("2.9 REST API", "REST API menyediakan endpoint HTTP seperti GET, POST, PUT, PATCH, dan DELETE untuk komunikasi antara frontend dan backend."),
        ("2.10 Database PostgreSQL", "Project ini memakai PostgreSQL. Data yang disimpan meliputi user, pengaduan, layanan administrasi, kategori layanan, form dinamis, lampiran, notifikasi, dan riwayat status."),
        ("2.11 Deployment", "Deployment dapat dilakukan di Railway dengan PostgreSQL managed database. Aplikasi membaca konfigurasi database dari environment variable."),
    ]
    for title, body in sections:
        add_heading(doc, title, 2 if title.count(".") == 1 else 3)
        add_justified(doc, body)


def add_chapter_3(doc):
    add_heading(doc, "BAB III PERANCANGAN SISTEM", 1)
    add_heading(doc, "3.1 Studi Kasus", 2)
    add_justified(
        doc,
        "Studi kasus yang diambil adalah sistem pelayanan masyarakat berbasis web. "
        "Permasalahan utama adalah pengaduan dan layanan administrasi yang masih "
        "manual, sehingga dibutuhkan platform terpadu untuk warga dan admin."
    )
    add_heading(doc, "3.2 Fitur Aplikasi", 2)
    add_numbered(doc, [
        "Autentikasi: login dan registrasi pengguna dengan role Admin, Warga, dan Superadmin.",
        "Pengaduan masyarakat: warga dapat mengirim keluhan, kritik, atau laporan masalah.",
        "Layanan administrasi: warga dapat mengajukan layanan berdasarkan kategori.",
        "Form dinamis: admin dapat mengatur field isian untuk setiap kategori layanan.",
        "Status pelayanan: admin memperbarui status, warga memantau status secara real-time.",
        "Riwayat status: setiap perubahan status dicatat sebagai history log.",
        "Lampiran file: warga dapat mengunggah dokumen atau bukti pendukung.",
        "Notifikasi: user menerima informasi perubahan status.",
        "PDF laporan: detail pengaduan dan layanan dapat diunduh sebagai PDF.",
    ])
    add_heading(doc, "3.3 Desain Database", 2)
    add_heading(doc, "3.3.1 ERD (Entity Relationship Diagram)", 3)
    add_image_or_placeholder(doc, find_image("docx_extracted/word/media/image2.jpg", "screenshots/image2.jpg"), "Gambar 3.3.1 ERD Sistem Pelayanan Masyarakat")
    add_heading(doc, "3.3.2 TRD (Table Relationship Diagram)", 3)
    add_image_or_placeholder(doc, find_image("docx_extracted/word/media/image3.png", "screenshots/image3.png"), "Gambar 3.3.2 TRD Sistem Pelayanan Masyarakat")
    add_heading(doc, "3.3.3 Spesifikasi Tabel Database", 3)
    add_table(doc, ["Tabel", "Fungsi", "Kolom Utama"], [
        ["tb_user", "Menyimpan data pengguna", "id, nama_lengkap, email, password, role, status_aktif"],
        ["tb_pengaduan", "Menyimpan pengaduan warga", "id, judul, deskripsi, status, prioritas, user_id"],
        ["tb_layanan_administrasi", "Menyimpan permohonan layanan", "id, nomor_permohonan, status, user_id, kategori_id"],
        ["tb_kategori_layanan", "Menyimpan kategori layanan", "id, nama_kategori, deskripsi, is_active"],
        ["tb_form_field", "Menyimpan konfigurasi field form", "id, label, tipe, required, kategori_id"],
        ["tb_jawaban_form", "Menyimpan jawaban form layanan", "id, layanan_id, field_id, nilai"],
        ["tb_lampiran_file", "Menyimpan lampiran", "id, nama_file, nama_tersimpan, referensi_id"],
        ["tb_notifikasi", "Menyimpan notifikasi user", "id, judul, pesan, user_id, sudah_dibaca"],
        ["tb_riwayat_status", "Menyimpan riwayat status pengaduan", "id, referensi_id, status_lama, status_baru"],
        ["tb_riwayat_status_layanan", "Menyimpan riwayat status layanan", "id, layanan_id, status_lama, status_baru"],
    ])
    add_heading(doc, "3.4 Desain OOP & Rancangan API", 2)
    add_heading(doc, "3.4.1 Class Diagram", 3)
    add_image_or_placeholder(doc, find_image("docx_extracted/word/media/image4.jpg", "screenshots/image4.jpg"), "Gambar 3.4.1 Class Diagram SmartPelayanan")
    add_heading(doc, "3.4.2 Hubungan antar Class", 3)
    add_bullets(doc, [
        "BaseEntity diwarisi entity utama untuk id, createdAt, dan updatedAt.",
        "User berelasi dengan Pengaduan dan LayananAdministrasi.",
        "KategoriLayanan berelasi dengan LayananAdministrasi dan FormField.",
        "LayananAdministrasi berelasi dengan JawabanForm dan RiwayatStatusLayanan.",
        "Pengaduan berelasi dengan ResponPengaduan dan RiwayatStatus.",
    ])
    add_heading(doc, "3.4.3 Identifikasi Layer Sistem", 3)
    add_bullets(doc, [
        "@Entity merepresentasikan tabel database.",
        "@Repository menangani akses data.",
        "@Service menangani logika bisnis.",
        "@RestController menangani request dan response API.",
        "Template HTML menangani tampilan halaman login, admin, warga, dan superadmin.",
    ])
    add_heading(doc, "3.5 Rancangan REST API", 2)
    add_image_or_placeholder(doc, find_image("docx_extracted/word/media/image5.jpg", "screenshots/image5.jpg"), "Gambar 3.5.1 Diagram Endpoint REST API")
    add_table(doc, ["Modul", "Endpoint", "Keterangan"], [
        ["Auth", "/api/v1/auth/login", "Login user dan admin"],
        ["Auth", "/api/v1/auth/register", "Registrasi user"],
        ["Pengaduan", "/api/v1/pengaduan/{id}/detail", "Detail pengaduan dan riwayat status"],
        ["Admin Pengaduan", "/api/v1/admin/pengaduan/{id}/status", "Update status pengaduan"],
        ["Warga Pengaduan", "/api/v1/warga/pengaduan", "Tambah dan lihat pengaduan warga"],
        ["Layanan", "/api/v1/warga/layanan", "Tambah dan lihat layanan warga"],
        ["Admin Layanan", "/api/v1/admin/layanan/{id}/status", "Update status layanan"],
        ["Laporan", "/api/v1/laporan/pengaduan/{id}/pdf", "Unduh PDF pengaduan"],
        ["Laporan", "/api/v1/laporan/layanan/{id}/pdf", "Unduh PDF layanan"],
        ["Notifikasi", "/api/v1/notifikasi", "Daftar notifikasi user"],
    ])


def add_chapter_4(doc):
    add_heading(doc, "BAB IV IMPLEMENTASI SISTEM", 1)
    sections = [
        ("4.1 Instalasi dan Konfigurasi", "Aplikasi menggunakan Java 17, Spring Boot 3.2.5, PostgreSQL, Maven, dan Flyway. Konfigurasi Railway dibaca melalui environment variable database."),
        ("4.2 Implementasi Entity dan Repository", "Entity dibuat pada package com.smartpelayanan.entity. Repository dibuat pada package com.smartpelayanan.repository dengan JpaRepository."),
        ("4.3 Implementasi Service Layer", "PengaduanService dan PengaduanServiceImpl menangani operasi pengaduan dan layanan agar controller tidak menampung semua logika."),
        ("4.4 Implementasi Controller Layer", "Controller menerima HTTP request, memvalidasi token, memanggil repository/service, lalu mengembalikan response JSON atau PDF."),
        ("4.5 Implementasi REST API", "REST API digunakan oleh halaman frontend untuk login, mengambil data dashboard, membuat pengaduan, membuat layanan, update status, notifikasi, dan unduh PDF."),
        ("4.6 Implementasi Layered Architecture", "Layer aplikasi dipisah menjadi Controller, Service, Repository, Entity, dan Database agar kode lebih rapi."),
        ("4.7 Implementasi Exception Handling", "Controller memakai try-catch dan response error agar user mendapat pesan ketika request gagal."),
        ("4.8 Hasil Implementasi Sistem", "Sistem sudah memiliki login, registrasi, dashboard admin/warga, pengaduan, layanan administrasi, form dinamis, riwayat status, notifikasi, upload lampiran, dan PDF laporan."),
    ]
    for title, body in sections:
        add_heading(doc, title, 2)
        add_justified(doc, body)


def add_chapter_5(doc):
    add_heading(doc, "BAB V DEPLOYMENT DAN CLOUDFLARE TUNNEL", 1)
    add_bullets(doc, [
        "Tujuan deployment adalah membuat aplikasi dapat diakses secara online.",
        "Railway dapat digunakan untuk menjalankan aplikasi Spring Boot dan PostgreSQL.",
        "Database PostgreSQL dibuat melalui Railway PostgreSQL service.",
        "Aplikasi membaca koneksi database dari environment variable.",
        "Flyway menjalankan migrasi schema saat aplikasi start.",
        "Cloudflare Tunnel dapat digunakan bila aplikasi dijalankan dari server lokal.",
    ])
    add_heading(doc, "5.7 URL Publik Sistem", 2)
    add_justified(doc, "URL publik sistem dapat diisi setelah aplikasi berhasil dideploy, misalnya melalui Railway atau Cloudflare Tunnel.")
    add_heading(doc, "5.8 Hasil Deployment", 2)
    add_image_or_placeholder(doc, find_image("screenshots/deployment_browser.png", "screenshots/18_Hasil_Deployment.png"), "Gambar 5.8.1 Hasil Deployment di Browser")
    add_heading(doc, "5.9 Evaluasi Deployment", 2)
    add_bullets(doc, [
        "Aplikasi dapat berjalan di lingkungan cloud.",
        "Database menggunakan PostgreSQL managed service.",
        "Migrasi database berjalan otomatis.",
        "Konfigurasi sensitif seperti password database dan JWT secret disimpan melalui environment variable.",
    ])


def add_chapter_6(doc):
    add_heading(doc, "BAB VI PENGUJIAN SISTEM", 1)
    tests = [
        ("6.1 Tujuan Pengujian", "Pengujian dilakukan untuk memastikan seluruh fitur berjalan sesuai kebutuhan."),
        ("6.2 Pengujian Login dan Registrasi", "User dapat registrasi dan login sesuai role."),
        ("6.3 Pengujian Data Layanan", "Warga dapat mengajukan layanan dan admin dapat memproses status."),
        ("6.4 Pengujian Kategori dan Form Dinamis", "Admin dapat mengatur kategori dan field form."),
        ("6.5 Pengujian Upload Lampiran", "File bukti dan dokumen pendukung dapat diunggah."),
        ("6.6 Pengujian Permohonan Layanan", "Permohonan layanan dapat dibuat, dilihat, dan diubah statusnya."),
        ("6.7 Pengujian Pengaduan", "Pengaduan dapat dikirim dan dikelola admin."),
        ("6.8 Pengujian Riwayat", "Perubahan status tercatat di riwayat detail dan PDF."),
    ]
    for title, body in tests:
        add_heading(doc, title, 2)
        add_justified(doc, body)
    add_image_or_placeholder(doc, find_image("screenshots/postman_result.png", "screenshots/17_Hasil_Postman.png"), "Gambar 6.8.1 Hasil Pengujian Endpoint API")


def add_chapter_7(doc):
    add_heading(doc, "BAB VII PENUTUP", 1)
    add_heading(doc, "7.1 Kesimpulan", 2)
    add_bullets(doc, [
        "SmartPelayanan berhasil dirancang dan diimplementasikan sebagai sistem pelayanan masyarakat berbasis web.",
        "Aplikasi menggunakan Java, Spring Boot, Spring Data JPA, REST API, PostgreSQL, dan Flyway.",
        "Fitur utama meliputi autentikasi, pengaduan, layanan administrasi, form dinamis, notifikasi, riwayat status, upload lampiran, dan PDF laporan.",
        "Konsep OOP diterapkan melalui entity, inheritance BaseEntity, encapsulation, abstraction repository/service, dan polymorphism melalui interface Spring Data.",
    ])
    add_heading(doc, "7.2 Saran", 2)
    add_bullets(doc, [
        "Menambahkan dashboard analitik yang lebih detail.",
        "Mengembangkan aplikasi mobile Android/iOS.",
        "Menambahkan fitur audit log yang lebih lengkap.",
        "Meningkatkan keamanan dengan refresh token dan konfigurasi production profile.",
    ])


def add_screenshots(doc):
    add_heading(doc, "DAFTAR FOTO DAN SCREENSHOT", 1)
    images = [
        ("screenshots/1_Login_Page.png", "Gambar 1 Halaman Login Aplikasi - URL http://localhost:8082/login"),
        ("screenshots/2_Register_Page.png", "Gambar 2 Halaman Registrasi Aplikasi - URL http://localhost:8082/register"),
        ("screenshots/4_Admin_Dashboard.png", "Gambar 3 Dashboard Admin - URL http://localhost:8082/admin/dashboard"),
        ("screenshots/7_Warga_Dashboard_LoggedIn.png", "Gambar 4 Dashboard Warga - URL http://localhost:8082/warga/dashboard"),
        ("screenshots/8_Warga_Pengaduan.png", "Gambar 5 Form Pengaduan Masyarakat"),
        ("screenshots/3_Public_Pengaduan.png", "Gambar 6 Pengaduan Publik Tanpa Login"),
        ("screenshots/5_Superadmin_Dashboard.png", "Gambar 7 Dashboard Superadmin"),
        ("screenshots/layanan_form.png", "Gambar 8 Form Permohonan Layanan"),
        ("screenshots/status_permohonan.png", "Gambar 9 Status Permohonan"),
        ("screenshots/riwayat_pelayanan.png", "Gambar 10 Riwayat Pelayanan"),
        ("screenshots/postman_login.png", "Gambar 11 Pengujian Login di Postman"),
        ("screenshots/postman_pengaduan.png", "Gambar 12 Pengujian Pengaduan di Postman"),
        ("screenshots/postman_layanan.png", "Gambar 13 Pengujian Layanan di Postman"),
        ("screenshots/railway_deploy.png", "Gambar 14 Hasil Deployment Railway"),
        ("screenshots/database_postgresql.png", "Gambar 15 Database PostgreSQL"),
    ]
    for rel_path, caption in images:
        add_image_or_placeholder(doc, find_image(rel_path), caption)


def add_code_oop_explanation(doc):
    add_heading(doc, "PENJELASAN KODE, ALUR, DAN PENERAPAN OOP", 1)
    add_heading(doc, "A. Struktur Kode Project", 2)
    add_table(doc, ["Package/File", "Fungsi"], [
        ["com.smartpelayanan.entity", "Berisi class entity yang dipetakan ke tabel PostgreSQL."],
        ["com.smartpelayanan.repository", "Berisi interface repository untuk CRUD dan query database."],
        ["com.smartpelayanan.service", "Berisi kontrak service untuk logika bisnis."],
        ["com.smartpelayanan.service.impl", "Berisi implementasi service."],
        ["com.smartpelayanan.controller", "Berisi REST controller dan page controller."],
        ["templates/*.html", "Berisi tampilan login, register, dashboard admin, warga, dan superadmin."],
        ["db/migration", "Berisi migrasi Flyway untuk membuat schema PostgreSQL."],
    ])
    add_heading(doc, "B. Alur Login", 2)
    add_numbered(doc, [
        "User membuka halaman /login.",
        "User memasukkan email dan password.",
        "Frontend mengirim request POST ke /api/v1/auth/login.",
        "AuthController memvalidasi data login.",
        "Jika valid, JwtUtils membuat token JWT.",
        "Token disimpan di localStorage dan user diarahkan ke dashboard sesuai role.",
    ])
    add_heading(doc, "C. Alur Pengaduan Warga", 2)
    add_numbered(doc, [
        "Warga membuka dashboard dan mengisi form pengaduan.",
        "Frontend mengirim data ke /api/v1/warga/pengaduan.",
        "WebController membuat object Pengaduan dan menyimpan ke tb_pengaduan.",
        "Admin menerima notifikasi pengaduan baru.",
        "Admin dapat memperbarui status pengaduan.",
        "Setiap perubahan status disimpan ke tb_riwayat_status.",
        "Warga dapat melihat detail, riwayat status, dan mengunduh PDF.",
    ])
    add_heading(doc, "D. Alur Layanan Administrasi", 2)
    add_numbered(doc, [
        "Admin membuat kategori layanan dan field form dinamis.",
        "Warga memilih kategori layanan dan mengisi form.",
        "Jawaban form disimpan ke tb_jawaban_form.",
        "Admin memproses permohonan dan mengubah status.",
        "Perubahan status disimpan ke tb_riwayat_status_layanan.",
        "Warga dapat memantau status dan mengunduh PDF laporan layanan.",
    ])
    add_heading(doc, "E. Penerapan Encapsulation", 2)
    add_justified(
        doc,
        "Encapsulation terlihat pada entity seperti User, Pengaduan, dan LayananAdministrasi. "
        "Atribut dibuat private, lalu diakses memakai getter dan setter. Dengan cara ini, data tidak "
        "diubah langsung dari luar class."
    )
    add_heading(doc, "F. Penerapan Inheritance", 2)
    add_justified(
        doc,
        "Class BaseEntity menjadi parent untuk entity utama. Entity yang extends BaseEntity otomatis "
        "memiliki id, createdAt, dan updatedAt. Ini mengurangi duplikasi kode."
    )
    add_heading(doc, "G. Penerapan Abstraction", 2)
    add_justified(
        doc,
        "Repository seperti UserRepository dan PengaduanRepository adalah abstraction. Controller tidak "
        "perlu menulis SQL langsung; cukup memanggil method repository seperti findByEmail atau findByUserId."
    )
    add_heading(doc, "H. Penerapan Polymorphism", 2)
    add_justified(
        doc,
        "Polymorphism terlihat pada penggunaan JpaRepository untuk banyak entity. Bentuk operasinya sama "
        "seperti save, findById, findAll, dan deleteById, tetapi objek yang diproses berbeda-beda."
    )
    add_heading(doc, "I. Penerapan Exception Handling", 2)
    add_justified(
        doc,
        "Controller menggunakan try-catch untuk menangani error. Contohnya ketika data tidak ditemukan, "
        "controller mengembalikan response error agar frontend dapat menampilkan pesan gagal."
    )
    add_heading(doc, "J. Implementasi Migrasi PostgreSQL", 2)
    add_justified(
        doc,
        "Flyway membaca file SQL pada src/main/resources/db/migration. Saat aplikasi dijalankan di Railway, "
        "Flyway membuat tabel yang diperlukan sebelum Hibernate melakukan validasi entity."
    )


def build_document():
    doc = setup_document()
    add_cover(doc)
    add_preface(doc)
    add_table_of_contents(doc)
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    doc.add_page_break()
    add_chapter_7(doc)
    doc.add_page_break()
    add_code_oop_explanation(doc)
    doc.add_page_break()
    add_screenshots(doc)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_document()
    print(f"Laporan berhasil dibuat: {output}")
